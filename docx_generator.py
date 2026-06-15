import os, copy, tempfile, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CUENTAS_LABELS = {
    'vrtcl': 'VRTCL',
    'correo': 'Correo institucional',
    'sap': 'SAP',
    'pos': 'POS',
    'req': 'Portal de Requisición'
}

def fmt_fecha(d):
    if not d: return '—'
    try:
        p = d.split('-')
        return f"{p[2]}/{p[1]}/{p[0]}"
    except:
        return d

def _set_cell_color(cell, hex_color):
    """Aplica color de fondo a una celda SIN tocar el texto."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Eliminar shd previo para no acumular
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def _set_cell_borders(cell, color='BBBBBB'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcB = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcB.append(b)
    tcPr.append(tcB)

def _run(para, text, bold=False, size=9, italic=False, color_hex=None):
    r = para.add_run(str(text))
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color_hex:
        r.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    return r

def _heading_row(table, text, cols, bg='1A1815', fg='FFFFFF'):
    """Fila de encabezado oscura con texto blanco."""
    row = table.add_row()
    cell = row.cells[0]
    # Merge si hay varias columnas
    if cols > 1:
        for i in range(1, cols):
            cell = cell.merge(row.cells[i])
    _set_cell_color(cell, bg)
    _set_cell_borders(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(p, text, bold=True, size=9, color_hex=fg)
    return row

def _table_header(table, headers, widths_cm, bg='E8E8E8'):
    """Fila de cabecera de tabla."""
    row = table.add_row()
    for i, (hdr, w) in enumerate(zip(headers, widths_cm)):
        cell = row.cells[i]
        cell.width = Cm(w)
        _set_cell_color(cell, bg)
        _set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, hdr, bold=True, size=8)

def _data_row(table, values, bg='FFFFFF', sizes=None):
    sizes = sizes or [9]*len(values)
    row = table.add_row()
    for i, (val, sz) in enumerate(zip(values, sizes)):
        cell = row.cells[i]
        _set_cell_color(cell, bg)
        _set_cell_borders(cell)
        _run(cell.paragraphs[0], val or '—', size=sz)
    return row

def build_doc(emp, equipos_asignados, tipo='alta'):
    base_dir  = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(base_dir, 'static', 'logo.png')
    plantilla = os.path.join(base_dir, 'plantilla', 'FORMATO_DE_RESGUARDO.docx')

    # ── Crear documento limpio (márgenes carta) ──
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)
        sec.page_width    = Cm(21.59)
        sec.page_height   = Cm(27.94)

    # ── 1. HEADER ──────────────────────────────────────────────
    ht = doc.add_table(rows=1, cols=2)
    ht.style = 'Table Grid'
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Celda logo — fondo blanco
    cl = ht.rows[0].cells[0]
    cl.width = Cm(3.2)
    _set_cell_color(cl, 'FFFFFF')
    _set_cell_borders(cl, 'CCCCCC')
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        pl.add_run().add_picture(logo_path, width=Cm(2.5))

    # Celda título — fondo oscuro, texto blanco
    ct = ht.rows[0].cells[1]
    _set_cell_color(ct, '1F2937')
    _set_cell_borders(ct, '1F2937')
    pt = ct.paragraphs[0]
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pt, 'RESGUARDO DE EQUIPO INFORMÁTICO', bold=True, size=13, color_hex='FFFFFF')
    p2 = ct.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, 'HOSPITAL ESCANDÓN — ÁREA DE SISTEMAS', bold=True, size=10, color_hex='D1D5DB')

    # Subtítulo tipo
    p3 = ct.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if tipo == 'alta':
        _run(p3, '▶  ENTREGA DE EQUIPO Y CUENTAS  ◀', bold=True, size=8.5, color_hex='FCD34D')
    else:
        _run(p3, '▶  DEVOLUCIÓN DE EQUIPO  ◀', bold=True, size=8.5, color_hex='FCA5A5')

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 2. INFO GENERAL ────────────────────────────────────────
    it = doc.add_table(rows=4, cols=4)
    it.style = 'Table Grid'

    fecha_lbl = 'Fecha de Alta:' if tipo == 'alta' else 'Fecha de Baja:'
    fecha_val = fmt_fecha(emp.get('fecha_alta')) if tipo == 'alta' else fmt_fecha(emp.get('fecha_baja'))

    data = [
        (fecha_lbl,             fecha_val,
         'Dirección/Subdirección:', emp.get('direccion') or '—'),
        ('Responsable:',        emp.get('nombre','—'),
         'Departamento/Área:',  emp.get('area') or '—'),
        ('Ubicación:',          emp.get('ubicacion') or '—',
         'Fecha de Asignación:', fmt_fecha(emp.get('fecha_alta'))),
        ('Teléfono:',           emp.get('telefono') or '—',
         '', ''),
    ]

    for ri, (l1,v1,l2,v2) in enumerate(data):
        row = it.rows[ri]
        for ci, txt in enumerate([l1,v1,l2,v2]):
            cell = row.cells[ci]
            _set_cell_color(cell, 'F9FAFB' if ci%2==0 else 'FFFFFF')
            _set_cell_borders(cell, 'D1D5DB')
            p = cell.paragraphs[0]
            _run(p, txt, bold=(ci%2==0), size=8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # ── 3. TABLA EQUIPOS ───────────────────────────────────────
    et = doc.add_table(rows=0, cols=4)
    et.style = 'Table Grid'
    # Heading
    _heading_row(et, '  DESCRIPCIÓN DE EQUIPO ' + ('ASIGNADO' if tipo=='alta' else 'A DEVOLVER'), 4)
    _table_header(et, ['Tipo / Categoría','Descripción y Especificaciones','No. de Serie / IP','Cant.'],
                  [3.5, 9.0, 4.0, 1.0])

    equipos = equipos_asignados or []
    if equipos:
        for i, eq in enumerate(equipos):
            bg = 'FFFFFF' if i%2==0 else 'F3F4F6'
            desc = f"{eq.get('marca','')} {eq.get('modelo','')}".strip()
            specs = []
            if eq.get('procesador'): specs.append(f"CPU: {eq['procesador']}")
            if eq.get('ram'):        specs.append(f"RAM: {eq['ram']}")
            if eq.get('almacenamiento'): specs.append(eq['almacenamiento'])
            if specs: desc += f"  [{', '.join(specs)}]"
            ns = eq.get('num_serie','') or eq.get('ip_asignada','') or '—'
            _data_row(et, [eq.get('tipo',''), desc or '—', ns, str(eq.get('cantidad',1))],
                      bg=bg, sizes=[8.5,8.5,8.5,8.5])
    else:
        row = et.add_row()
        cell = row.cells[0].merge(row.cells[3])
        _set_cell_color(cell, 'FFF7ED')
        _run(cell.paragraphs[0], 'Sin equipos asignados', italic=True, size=8.5, color_hex='9CA3AF')

    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # ── 4. TABLA CUENTAS ───────────────────────────────────────
    cuentas = emp.get('cuentas', {})
    if isinstance(cuentas, str):
        cuentas = json.loads(cuentas)
    activas = {k:v for k,v in cuentas.items() if v}

    ct2 = doc.add_table(rows=0, cols=3)
    ct2.style = 'Table Grid'
    _heading_row(ct2, '  CUENTAS Y ACCESOS ' + ('ASIGNADOS' if tipo=='alta' else 'A DAR DE BAJA'), 3)
    _table_header(ct2, ['Sistema / Plataforma','Usuario / Cuenta','Acción'],
                  [4.0, 8.5, 5.0])

    if activas:
        accion_txt = 'Se asigna acceso' if tipo=='alta' else 'Se da de baja'
        for i,(k,v) in enumerate(activas.items()):
            bg = 'FFFFFF' if i%2==0 else 'F3F4F6'
            _data_row(ct2, [CUENTAS_LABELS.get(k,k.upper()), v, accion_txt], bg=bg, sizes=[8.5,8.5,8.5])
    else:
        row = ct2.add_row()
        cell = row.cells[0].merge(row.cells[2])
        _set_cell_color(cell, 'FFF7ED')
        _run(cell.paragraphs[0], 'Sin cuentas asignadas', italic=True, size=8.5, color_hex='9CA3AF')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 5. TEXTO LEGAL ─────────────────────────────────────────
    if tipo == 'alta':
        parrafos = [
            'Hago constar que recibo conforme el equipo y accesorios descritos, propiedad de HOSPITAL ESCANDÓN, '
            'mismos que quedan bajo mi responsabilidad a partir de esta fecha.',
            'Me comprometo a utilizarlos únicamente para actividades institucionales, a conservarlos en buen '
            'estado considerando el desgaste normal por uso, y a no instalar software distinto al autorizado '
            'por la Institución, en cumplimiento con la Ley Federal de Derechos de Autor.',
            'Asimismo, me hago responsable del resguardo de los accesos y cuentas institucionales que me han '
            'sido asignados, comprometiéndome a no compartirlos y a usarlos exclusivamente para fines laborales.'
        ]
    else:
        parrafos = [
            'Hago constar que entrego la clave, equipo y accesorios descritos, propiedad de HOSPITAL ESCANDÓN, '
            'que se encontraban bajo mi resguardo, mismos que se devuelven en las condiciones en que fueron '
            'asignados, considerando únicamente el desgaste normal por uso.',
            'Manifiesto que el equipo se entrega sin información personal o ajena a las actividades '
            'institucionales y sin software distinto al autorizado por la Institución, en cumplimiento con '
            'la Ley Federal de Derechos de Autor.',
            'Con la presente devolución, quedo liberado(a) de las responsabilidades administrativas y legales '
            'derivadas del resguardo, a partir de la fecha de recepción por el área responsable.'
        ]

    for txt in parrafos:
        p = doc.add_paragraph()
        _run(p, txt, size=8.5)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 6. FIRMAS ──────────────────────────────────────────────
    ft = doc.add_table(rows=0, cols=2)
    ft.style = 'Table Grid'
    # Líneas de firma
    r1 = ft.add_row()
    for cell in r1.cells:
        _set_cell_color(cell, 'FFFFFF')
        _set_cell_borders(cell, 'FFFFFF')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, '_' * 40, size=9)
    # Títulos
    r2 = ft.add_row()
    titulos = ['ÁREA DE SISTEMAS',
               'RECIBÍ / CONFORME' if tipo=='alta' else 'ENTREGUÉ / CONFORME']
    for cell, txt in zip(r2.cells, titulos):
        _set_cell_color(cell, 'F3F4F6')
        _set_cell_borders(cell, 'D1D5DB')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, txt, bold=True, size=8.5)
    # Nombres
    r3 = ft.add_row()
    nombres = ['Firma y sello', emp.get('nombre','')]
    for cell, txt in zip(r3.cells, nombres):
        _set_cell_color(cell, 'FFFFFF')
        _set_cell_borders(cell, 'D1D5DB')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, txt, size=8.5, italic=(txt=='Firma y sello'))

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 7. PIE ─────────────────────────────────────────────────
    pp = doc.add_paragraph()
    _run(pp, 'Este resguardo solo será VÁLIDO si tiene firma y/o sello del RESPONSABLE DE CONTROL DE EQUIPOS '
             'Y ACCESORIOS INFORMÁTICOS AUTORIZADO. — Este documento cancela y sustituye cualquier resguardo '
             'anterior con datos iguales en No. de Serie, Modelo o descripción del equipo.',
         size=7, italic=True, color_hex='6B7280')

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    tmp.close()
    return tmp.name

def generar_docx_alta(emp, equipos): return build_doc(emp, equipos, 'alta')
def generar_docx_baja(emp, equipos): return build_doc(emp, equipos, 'baja')

# ── 8. NUEVA FUNCIÓN PARA PRÉSTAMOS TEMPORALES ─────────────────
def generar_docx_prestamo(prestamo, equipos):
    base_dir  = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(base_dir, 'static', 'logo.png')

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)
        sec.page_width    = Cm(21.59)
        sec.page_height   = Cm(27.94)

    # ── 1. HEADER ──
    ht = doc.add_table(rows=1, cols=2)
    ht.style = 'Table Grid'
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER

    cl = ht.rows[0].cells[0]
    cl.width = Cm(3.2)
    _set_cell_color(cl, 'FFFFFF')
    _set_cell_borders(cl, 'CCCCCC')
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        pl.add_run().add_picture(logo_path, width=Cm(2.5))

    ct = ht.rows[0].cells[1]
    _set_cell_color(ct, '1F2937')
    _set_cell_borders(ct, '1F2937')
    pt = ct.paragraphs[0]
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pt, 'RESGUARDO TEMPORAL DE EQUIPO', bold=True, size=13, color_hex='FFFFFF')
    p2 = ct.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, 'HOSPITAL ESCANDÓN — ÁREA DE SISTEMAS', bold=True, size=10, color_hex='D1D5DB')

    p3 = ct.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p3, '▶  PRÉSTAMO TEMPORAL  ◀', bold=True, size=8.5, color_hex='60A5FA')

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 2. INFO GENERAL ──
    it = doc.add_table(rows=3, cols=4)
    it.style = 'Table Grid'

    data = [
        ('Fecha de Préstamo:', fmt_fecha(prestamo.get('fecha_prestamo')),
         'Devolución Esperada:', fmt_fecha(prestamo.get('fecha_esperada'))),
        ('Solicitante:', prestamo.get('nombre_solicitante', '—'),
         'Área/Departamento:', prestamo.get('area', '—')),
        ('Motivo/Notas:', prestamo.get('notas', '—'),
         'Estado del Trámite:', prestamo.get('estado', '—').upper()),
    ]

    for ri, (l1,v1,l2,v2) in enumerate(data):
        row = it.rows[ri]
        for ci, txt in enumerate([l1,v1,l2,v2]):
            cell = row.cells[ci]
            _set_cell_color(cell, 'F9FAFB' if ci%2==0 else 'FFFFFF')
            _set_cell_borders(cell, 'D1D5DB')
            p = cell.paragraphs[0]
            _run(p, txt, bold=(ci%2==0), size=8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # ── 3. TABLA EQUIPOS ──
    et = doc.add_table(rows=0, cols=4)
    et.style = 'Table Grid'
    _heading_row(et, '  EQUIPOS OTORGADOS EN PRÉSTAMO Y ESTADO DE ENTREGA', 4)
    _table_header(et, ['Tipo Equipo', 'Descripción / Marca / Serie', 'Condición / Notas de Entrega', 'Cant.'],
                  [3.5, 6.5, 6.5, 1.0])

    if equipos:
        for i, eq in enumerate(equipos):
            bg = 'FFFFFF' if i%2==0 else 'F3F4F6'
            desc = f"{eq.get('marca','')} {eq.get('modelo','')}".strip()
            ns = eq.get('num_serie','') or '—'
            full_desc = f"{desc}\nS/N: {ns}"
            notas_est = eq.get('notas_estado', 'Buen estado')
            _data_row(et, [eq.get('tipo',''), full_desc, notas_est, str(eq.get('cantidad',1))],
                      bg=bg, sizes=[8.5,8.5,8.5,8.5])
    else:
        row = et.add_row()
        cell = row.cells[0].merge(row.cells[3])
        _set_cell_color(cell, 'FFF7ED')
        _run(cell.paragraphs[0], 'Sin equipos registrados', italic=True, size=8.5, color_hex='9CA3AF')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 4. TEXTO LEGAL ──
    parrafos = [
        'El solicitante reconoce recibir el/los equipo(s) en las condiciones descritas en el apartado anterior '
        'y se compromete a devolverlos en el mismo estado operativo y físico.',
        f'El equipo es otorgado bajo un esquema de préstamo temporal y deberá ser devuelto a más tardar el día '
        f'{fmt_fecha(prestamo.get("fecha_esperada"))} al Área de Sistemas.',
        'Cualquier daño, alteración de hardware/software, robo o extravío generado durante el periodo '
        'de préstamo será responsabilidad absoluta del solicitante.',
        'Queda estrictamente prohibida la instalación de software no autorizado o el almacenamiento de información '
        'personal en los equipos prestados.'
    ]

    for txt in parrafos:
        p = doc.add_paragraph()
        _run(p, txt, size=8.5)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 5. FIRMAS ──
    ft = doc.add_table(rows=0, cols=2)
    ft.style = 'Table Grid'
    r1 = ft.add_row()
    for cell in r1.cells:
        _set_cell_color(cell, 'FFFFFF')
        _set_cell_borders(cell, 'FFFFFF')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, '_' * 40, size=9)
    r2 = ft.add_row()
    titulos = ['ENTREGA (ÁREA DE SISTEMAS)', 'RECIBÍ / CONFORME (SOLICITANTE)']
    for cell, txt in zip(r2.cells, titulos):
        _set_cell_color(cell, 'F3F4F6')
        _set_cell_borders(cell, 'D1D5DB')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, txt, bold=True, size=8.5)
    r3 = ft.add_row()
    nombres = ['Firma y sello', prestamo.get('nombre_solicitante','')]
    for cell, txt in zip(r3.cells, nombres):
        _set_cell_color(cell, 'FFFFFF')
        _set_cell_borders(cell, 'D1D5DB')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, txt, size=8.5, italic=(txt=='Firma y sello'))

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 6. PIE ──
    pp = doc.add_paragraph()
    _run(pp, 'Este documento acredita el resguardo temporal de los activos fijos del Hospital Escandón. '
             'Su firma compromete al solicitante a la devolución íntegra en la fecha estipulada.',
         size=7, italic=True, color_hex='6B7280')

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    tmp.close()
    return tmp.name